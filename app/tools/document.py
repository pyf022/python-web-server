from pathlib import Path

import markdown as markdown_lib
from bs4 import BeautifulSoup
from docx import Document
from pdf2docx import Converter
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas

from app.storage.base import Storage
from app.tools.base import BaseTool, ToolExecutionError, ToolResult
from app.tools.utils import command_exists, require_field, run_command


class WordToPdfTool(BaseTool):
    name = "word_to_pdf"
    description = "Convert DOC/DOCX files to PDF with LibreOffice headless."
    input_schema = {
        "type": "object",
        "required": ["file_id"],
        "properties": {
            "file_id": {
                "type": "string",
                "description": "已上传的 Word 文件 ID，只支持 .doc 或 .docx。",
                "examples": ["a1b2c3.docx"],
            }
        },
    }
    timeout_seconds = 600

    def execute(self, input_data: dict, storage: Storage) -> ToolResult:
        file_id = require_field(input_data, "file_id")
        source = storage.path_for(file_id)
        if source.suffix.lower() not in {".doc", ".docx"}:
            raise ToolExecutionError("invalid_input", "word_to_pdf expects a .doc or .docx file.")
        if not command_exists("soffice"):
            raise ToolExecutionError("missing_dependency", "LibreOffice 'soffice' is required for Word to PDF conversion.")
        temp_dir = storage.make_temp_dir()
        run_command(
            ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(temp_dir), str(source)],
            timeout=self.timeout_seconds,
        )
        pdf_path = temp_dir / f"{source.stem}.pdf"
        if not pdf_path.exists():
            raise ToolExecutionError("conversion_failed", "LibreOffice did not produce a PDF output.")
        stored = storage.save_path(pdf_path, filename=f"{Path(storage.get(file_id)['filename']).stem}.pdf", content_type="application/pdf")
        return ToolResult(type="file", result_file_id=stored["file_id"], filename=stored["filename"], content_type=stored["content_type"])


class PdfToWordTool(BaseTool):
    name = "pdf_to_word"
    description = "Best-effort PDF to DOCX conversion. Scanned PDFs require a future OCR extension."
    input_schema = {
        "type": "object",
        "required": ["file_id"],
        "properties": {
            "file_id": {
                "type": "string",
                "description": "已上传的 PDF 文件 ID，只支持 .pdf。扫描版 PDF 属于尽力转换。",
                "examples": ["a1b2c3.pdf"],
            }
        },
    }
    timeout_seconds = 600

    def execute(self, input_data: dict, storage: Storage) -> ToolResult:
        file_id = require_field(input_data, "file_id")
        source = storage.path_for(file_id)
        if source.suffix.lower() != ".pdf":
            raise ToolExecutionError("invalid_input", "pdf_to_word expects a .pdf file.")
        temp_dir = storage.make_temp_dir()
        output = temp_dir / f"{source.stem}.docx"
        try:
            converter = Converter(str(source))
            converter.convert(str(output))
            converter.close()
        except Exception as exc:
            raise ToolExecutionError("conversion_failed", "PDF to Word conversion failed.", {"error": str(exc)}) from exc
        stored = storage.save_path(output, filename=f"{Path(storage.get(file_id)['filename']).stem}.docx", content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        return ToolResult(type="file", result_file_id=stored["file_id"], filename=stored["filename"], content_type=stored["content_type"])


class MarkdownToWordTool(BaseTool):
    name = "md_to_word"
    description = "Convert Markdown text or Markdown file to DOCX."
    input_schema = {
        "type": "object",
        "properties": {
            "markdown_text": {
                "type": "string",
                "description": "Markdown 文本。`markdown_text` 和 `file_id` 二选一。",
                "examples": ["# Title\n\nHello"],
            },
            "file_id": {
                "type": "string",
                "description": "已上传的 Markdown 文件 ID。`markdown_text` 和 `file_id` 二选一。",
                "examples": ["a1b2c3.md"],
            },
            "output_filename": {
                "type": "string",
                "description": "输出 Word 文件名。",
                "default": "document.docx",
                "examples": ["result.docx"],
            },
        },
    }

    def execute(self, input_data: dict, storage: Storage) -> ToolResult:
        markdown_text, base_name = _markdown_source(input_data, storage)
        output_filename = input_data.get("output_filename") or f"{base_name}.docx"
        temp_dir = storage.make_temp_dir()
        md_path = temp_dir / "input.md"
        md_path.write_text(markdown_text, encoding="utf-8")
        output = temp_dir / "output.docx"
        if command_exists("pandoc"):
            run_command(["pandoc", str(md_path), "-o", str(output)], timeout=self.timeout_seconds)
        else:
            _write_docx_fallback(markdown_text, output)
        stored = storage.save_path(output, filename=output_filename, content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
        return ToolResult(type="file", result_file_id=stored["file_id"], filename=stored["filename"], content_type=stored["content_type"])


class MarkdownToPdfTool(BaseTool):
    name = "md_to_pdf"
    description = "Convert Markdown text or Markdown file to PDF."
    input_schema = {
        "type": "object",
        "properties": {
            "markdown_text": {
                "type": "string",
                "description": "Markdown 文本。`markdown_text` 和 `file_id` 二选一。",
                "examples": ["# Title\n\nHello"],
            },
            "file_id": {
                "type": "string",
                "description": "已上传的 Markdown 文件 ID。`markdown_text` 和 `file_id` 二选一。",
                "examples": ["a1b2c3.md"],
            },
            "output_filename": {
                "type": "string",
                "description": "输出 PDF 文件名。",
                "default": "document.pdf",
                "examples": ["result.pdf"],
            },
        },
    }

    def execute(self, input_data: dict, storage: Storage) -> ToolResult:
        markdown_text, base_name = _markdown_source(input_data, storage)
        output_filename = input_data.get("output_filename") or f"{base_name}.pdf"
        temp_dir = storage.make_temp_dir()
        md_path = temp_dir / "input.md"
        md_path.write_text(markdown_text, encoding="utf-8")
        output = temp_dir / "output.pdf"
        if command_exists("pandoc"):
            run_command(["pandoc", str(md_path), "-o", str(output)], timeout=self.timeout_seconds)
        else:
            _write_pdf_fallback(markdown_text, output)
        stored = storage.save_path(output, filename=output_filename, content_type="application/pdf")
        return ToolResult(type="file", result_file_id=stored["file_id"], filename=stored["filename"], content_type=stored["content_type"])


def _markdown_source(input_data: dict, storage: Storage) -> tuple[str, str]:
    if input_data.get("markdown_text"):
        return str(input_data["markdown_text"]), "document"
    if input_data.get("file_id"):
        meta = storage.get(str(input_data["file_id"]))
        text = Path(meta["path"]).read_text(encoding=input_data.get("encoding", "utf-8"))
        return text, Path(meta["filename"]).stem
    raise ToolExecutionError("invalid_input", "Provide either markdown_text or file_id.")


def _write_docx_fallback(markdown_text: str, output: Path) -> None:
    document = Document()
    html = markdown_lib.markdown(markdown_text)
    soup = BeautifulSoup(html, "html.parser")
    for element in soup.find_all(["h1", "h2", "h3", "p", "li"]):
        text = element.get_text(" ", strip=True)
        if not text:
            continue
        if element.name == "h1":
            document.add_heading(text, level=1)
        elif element.name == "h2":
            document.add_heading(text, level=2)
        elif element.name == "h3":
            document.add_heading(text, level=3)
        elif element.name == "li":
            document.add_paragraph(text, style="List Bullet")
        else:
            document.add_paragraph(text)
    document.save(output)


def _write_pdf_fallback(markdown_text: str, output: Path) -> None:
    pdf = canvas.Canvas(str(output), pagesize=A4)
    width, height = A4
    y = height - 48
    soup = BeautifulSoup(markdown_lib.markdown(markdown_text), "html.parser")
    lines = [line for line in soup.get_text("\n").splitlines() if line.strip()]
    for line in lines:
        if y < 48:
            pdf.showPage()
            y = height - 48
        pdf.drawString(48, y, line[:110])
        y -= 18
    pdf.save()
